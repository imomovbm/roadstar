from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from django.conf import settings
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import mm
from reportlab.lib import colors
from django.utils.html import strip_tags
from django.utils.safestring import mark_safe
from io import BytesIO
from .models import ContractNew, ContractItem, ContractContext, Agreement, AgreementContext
from django.shortcuts import get_object_or_404
from decimal import Decimal
from django.http import HttpResponse
from reportlab.platypus import Image
import qrcode
import base64
import re
from django.core.mail import EmailMessage, get_connection
import requests
import time
import hashlib
from eskiz_sms import EskizSMS

def clean_html_text(text):
    if text:
        # Replace <div> and </div> and <p> with <br>
        text = re.sub(r'</?(div|p)>', '<br>', text, flags=re.IGNORECASE)
        # Remove double <br>
        text = re.sub(r'(<br\s*/?>)+', '<br>', text)
        # Remove leading/trailing <br>
        text = re.sub(r'^(<br\s*/?>)+', '', text)
        text = re.sub(r'(<br\s*/?>)+$', '', text)
        # Make sure all <br> are self-closing for XHTML/PDF
        text = text.replace('<br>', '<br/>')
    return text

def send_sms_ibnux(phone, message):
    eskiz = EskizSMS(email="magiccodeuz@gmail.com", password="0Hnwo2nRZzvjGF8d9wi5fWNYsV4kUgkeU55G3sMC")
    eskiz.send_sms(phone, message=message, from_whom='4546', callback_url=None)

def send_custom_email(subject, message, recipient_list, sender_email, sender_password):
    connection = get_connection(
        host='smtp.gmail.com',
        port=587,
        username=sender_email,
        password=sender_password,
        use_tls=True
    )

    email = EmailMessage(
        subject=subject,
        body=message,
        from_email=sender_email,
        to=recipient_list,
        connection=connection
    )
    email.send()

def send_custom_email_with_attachment(subject, body, recipient_list, sender_email, sender_password, pdf_file, filename):
    connection = get_connection(
        host='smtp.gmail.com',
        port=587,
        username=sender_email,
        password=sender_password,
        use_tls=True
    )

    email = EmailMessage(
        subject=subject,
        body=body,
        from_email=sender_email,
        to=recipient_list,
        connection=connection
    )

    email.attach(filename, pdf_file.getvalue(), "application/pdf")
    email.send(fail_silently=False)


def clean_paragraph_for_pdf(html):
    # Convert <b> to ReportLab font tag
    html = re.sub(r'<b>', "<font name='TimesNewRomanBold'>", html)
    html = re.sub(r'</b>', "</font>", html)
    # Normalize <br> to <br/> (self-closing)
    html = html.replace('<br>', '<br/>')
    return html

def get_qr_image_base64(pk):
    contract = get_object_or_404(ContractNew, pk=pk, is_public=True, deleted=False)
    url = f"https://roadstar.uz/contracts/public-pdf/{contract.public_token}/"
    if contract.canceled:
        url = f"https://roadstar.uz/contracts/canceled-pdf/{contract.public_token}/"
    qr = qrcode.make(url)
    buffer = BytesIO()
    qr.save(buffer, format="PNG")
    img_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
    return f"data:image/png;base64,{img_str}"

def get_qr_image(pk):
    contract = get_object_or_404(ContractNew, pk=pk, is_public=True, deleted=False)
    url = f"https://roadstar.uz/contracts/public-pdf/{contract.public_token}/"
    if contract.canceled:
        url = f"https://roadstar.uz/contracts/canceled-pdf/{contract.public_token}/"
    qr = qrcode.make(url)
    buffer = BytesIO()
    qr.save(buffer, format="PNG")
    buffer.seek(0)
    return Image(buffer, width=80, height=80)


def get_default_paragraph_text(paragraph_number):
    default_paragraphs = {
        2: {
            "header": { 1: "ШАРТНОМАНИНГ УМУМИЙ МИҚДОРИ ВА ХИСОБ-КИТОБ ШАРТЛАРИ."},
            "text": {
                        1:"<b>“СОТИБ ОЛУВЧИ”</b> қуйидаги тартибда ва муддатларда тўловларни амалга ошириш мажбуриятини олади: ҳар икки томон ушбу шартномани имзолаган кундан бошлаб 5 банк куни ичида шартноманинг умумий суммасининг 100 фоиз миқдорида олдиндан тўловни амалга оширади.<br/>",
                        2: "Агар <b>“СОТИБ ОЛУВЧИ”</b> шартноманинг 2.1-бандда назарда тутилган шартларни бажармаса,  ушбу шартномадан <b>“СОТУВЧИ”</b> етказиб берилган товарларнинг нархларини ўзгартириш ёки товарларни етказиб бериш бўйича ўз зиммасига олган мажбуриятларни бажаришдан бош тортиш ҳуқуқини ўзида сақлаб қолади.<br/>",
                        3: "Товарнинг 100 фоиз аванс маблағи сотувчининг ҳисоб рақамига келиб тушгандан сўнг 10 куни ичида етказиб берилади.<br/>",
                        4: "Товарлар шартноманинг 2.3-бандда кўрсатилган муддат ичида <b>“СОТИБ ОЛУВЧИ”</b> томонидан ўз кучи ва маблағи билан олиб кетилади.<br/>",
                    }
        },
        3: {
            "header": { 1: "ТОМОНЛАРНИНГ ЖАВОБГАРЛИГИ."},
            "text": {
            1: "<b>“СОТИБ ОЛУВЧИ”</b> томонидан шартноманинг 2.1-бандида назарда тутилган мажбуриятлар бажарилмаган ёки лозим даражада бажарилмаган тақдирда <b>“СОТИБ ОЛУВЧИ”</b> <b>“СОТУВЧИ”</b>га ҳар бир кечиктирилган кун учун тўлов суммасининг 0,4 фоизи миқдорида, лекин кечиктирилган тўлов суммасининг 50 (эллик) фоизидан кўп бўлмаган миқдорда пеня тўлайди.<br/>",
            2: "Товарларни танлаб олмаслик, шунингдек <b>“СОТУВЧИ”</b> шартномада белгиланган муддатда (даврда) уларни етказиб берганда <b>“СОТИБ ОЛУВЧИ”</b> товарларни олишни асоссиз равишда рад этганлик учун сотиб олувчи <b>“СОТУВЧИ”</b> га танлаб олинмаган (ўз муддатида олинмаган) товарлар қийматининг 10 фоизи миқдорида жарима тўлайди. Товарлар танлаб олинмаган (олиш асоссиз рад этилган) ҳолларда, <b>“СОТУВЧИ”</b> жарима ундиришдан ташқари, ушбу товарлар мавжудлигининг кафолатларини тақдим этган ҳолда, <b>“СОТИБ ОЛУВЧИ”</b> дан танлаб олинмаган (ўз муддатида олинмаган) товарлар қиймати тўланишини талаб қилишга ҳақлидир.<br/>",
            3: "Автомашиналарнинг бошқа давлатлар чегараларида навбатларда тўхтаб қолиши, турли эпидемиялар сабабли шахарларнинг ёпилиши ва бошқа ҳолатлар юзага келганда шартноманинг 2.4-бандида белгиланган етказиб бериш муддатининг ўтиши тўхтатилади.<br/>",
            4: "Агарда шартноманинг 2.3 бандида кўрсатилган муддат ичида <b>“СОТУВЧИ”</b> томонидан етказилган товарни <b>“СОТИБ ОЛУВЧИ”</b> қабул қилиб олмаса (асоссиз рад этса), ҳисобварақ фактурани асоссиз рад этса, ишончнома бермаса, юборилган товар-транспорт юк хатини асоссиз рад этса, <b>“СОТИБ ОЛУВЧИ”</b> томонидан олдиндан <b>“СОТУВЧИ”</b>нинг ҳисоб рақамига ўтказиб берилган маблағ қайтарилмайди ва <b>“СОТУВЧИ”</b> ҳисобида қолади.<br/>",
            }
        },
        4: {
            "header": { 1: "ФОРС-МАЖОР."},
            "text": {
            1: "Умумэътироф этилган форс-можор ҳолатлари юзага келган тақдирда, ҳеч бир томон мажбуриятларнинг тўлиқ ёки қисман бажарилмаганлиги учун жавобгар бўлмайди. Бундай мажбуриятларга эга бўлган томон бошқа томонни улар пайдо бўлган кундан бошлаб 2 (икки) календар кун ичида хабардор қилиши шарт.<br/>"
            }
        },
        5: {
            "header": { 1: "НИЗОЛАРНИ ҲАЛ ҚИЛИШ ТАРТИБИ ВА БОШҚА ШАРТЛАР"},
            "text": {
            1: "Мазкур шартнома бўйича томонлар ўртасида юзага келадиган низо ва келишмовчиликлар ўзаро музокаралар ўтказиш йўли билан ҳал этилади.<br/>",
            2: "Шартномани бажариш чоғида юзага келган низо ва келишмовчиликларни музокаралар йўли билан ҳал этишнинг имконияти бўлмаса, улар Ўзбекистон Республикасининг амалдаги қонунчилиги асосида Бухоро туманлараро иқтисодий судда ҳал этилади.<br/>",
            3: "Ушбу шартнома ҳар икки томон имзолаган кундан бошлаб кучга киради ва тўлиқ бажарилгунга қадар амал қилади.<br/>",
            4: "Шартномани ўзгартириш ёки муддатидан олдин бекор қилиш томонларнинг ўзаро розилиги билан амалга оширилиши ва ёзма равишда амалга оширилиши керак.<br/>",
            5: "Ушбу шартнома билан тартибга солинмаган масалалар Ўзбекистон Республикасининг амалдаги қонунчилиги билан тартибга солинади.<br/>",
            6: "Мазкур шартнома Ўзбекистон Республикасининг амалдаги қонунчилигига мувофиқ талқин этилади ва ижро этилади.<br/>",
            7: " Ушбу шартнома ҳар бири бир хил юридик кучга эга бўлган икки нусхада имзоланди.<br/>",
            }
        },
    }

    return default_paragraphs.get(paragraph_number, {
        "header": {1:None},
        "text": {1:"Paragraf matni topilmadi."}
    })

def get_paragraph(contract, paragraph_number):
    # Load all customizations for this contract
    custom_contexts = ContractContext.objects.filter(
        contract=contract,
        paragraph_number=paragraph_number,
        deleted=False
    )
    context_dict = {}
    for o in custom_contexts:
        context_dict[(o.key, o.value)] = o.paragraph_text

    # Load defaults from your function
    defaults = get_default_paragraph_text(paragraph_number)

    # Get header
    header = context_dict.get(('header', 1), defaults['header'][1])

    # Get all subparagraphs (with fallback to defaults)
    text_dict = {}
    for subkey, subval in defaults['text'].items():
        text_dict[subkey] = context_dict.get(('text', subkey), subval)

    return {'header': header, 'text': text_dict}

uNames = ["", "бир", "икки", "уч", "тўрт", "беш", "олти", "етти", "саккиз", "тўққиз"]
dNames = ["", "ўн", "йигирма", "ўттиз", "қирқ", "эллик", "олтмиш", "етмиш", "саксон", "тўқсон"]
cNames = ["", "бир юз", "икки юз", "уч юз", "тўрт юз", "беш юз", "олти юз", "етти юз", "саккиз юз", "тўққиз юз"]

def read_class(num_str):
    num_str = num_str.zfill(3)
    c, d, u = int(num_str[0]), int(num_str[1]), int(num_str[2])
    result = []
    if c: result.append(cNames[c])
    if d: result.append(dNames[d])
    if u: result.append(uNames[u])
    return " ".join(result).strip()


def num_to_uz_cyrillic_text(value):
    try:
        num = int(value)
    except (ValueError, TypeError):
        return value

    if num == 0:
        return "нол"

    parts = []
    num_str = str(num).zfill(12)
    billions = num_str[:-9]
    millions = num_str[-9:-6]
    thousands = num_str[-6:-3]
    rest = num_str[-3:]

    if int(billions) > 0:
        parts.append(read_class(billions) + " миллиард")
    if int(millions) > 0:
        parts.append(read_class(millions) + " миллион")
    if int(thousands) > 0:
        parts.append(read_class(thousands) + " минг")
    if int(rest) > 0:
        parts.append(read_class(rest))

    result = " ".join(parts)

    # Clean-ups
    result = result.replace("бир юз минг", "юз минг")
    result = result.replace("бир юз миллион", "юз миллион")
    result = result.replace("миллион минг", "миллион")
    result = result.replace("миллиард миллион", "миллиард")
    if result.strip() == "бир минг":
        result = "минг"
    return result.strip()

def spaced_float(value, decimals=2):
    try:
        value = float(value)
        parts = f"{value:,.{int(decimals)}f}".split(".")
        integer_part = parts[0].replace(",", " ")
        return f"{integer_part}.{parts[1]}" if len(parts) > 1 else integer_part
    except (ValueError, TypeError):
        return value

def generate_contract_pdf(pk):
    contract = get_object_or_404(ContractNew, pk=pk, deleted=False)

    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'TIMES.TTF')
    pdfmetrics.registerFont(TTFont('TimesNewRoman', font_path))
    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'TIMESBD.TTF')
    pdfmetrics.registerFont(TTFont('TimesNewRomanBold', font_path))



    sales = ContractItem.objects.filter(deleted=False, contract_id=contract.pk).exclude(item='1').select_related('item')

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=10*mm, rightMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY, fontName='TimesNewRoman', fontSize=8, leading=14))
    styles.add(ParagraphStyle(name='CenterBold', alignment=TA_CENTER, fontName='TimesNewRomanBold', fontSize=8))
    styles.add(ParagraphStyle(name='Centered', alignment=1, fontName='TimesNewRoman', fontSize=8, leading=12,))
    styles.add(ParagraphStyle(name='CenteredBold', parent=styles['Centered'], fontName='TimesNewRomanBold',))
    styles.add(ParagraphStyle(
        name='LeftBold',
        parent=styles['Normal'],
        fontName='TimesNewRomanBold',  # or 'TimesNewRomanBold' if registered
        alignment=0  # 0 = left
    ))
    styles.add(ParagraphStyle(
    name='RightBold',
    parent=styles['Normal'],
    fontName='TimesNewRomanBold',
    alignment=2  # 2 = right
    ))
    styles.add(ParagraphStyle(
    name='TinyCentered',
    parent=styles['Centered'],
    fontSize=7,       # or 8 depending on how small you want
    leading=9
    ))
    styles.add(ParagraphStyle(
        name='TinyCenteredBold',
        parent=styles['CenteredBold'],
        fontSize=7,
        leading=9
    ))
    elements = []

    # Title
    elements.append(Paragraph(f"ОЛДИ-СОТДИ ШАРТНОМАСИ", styles['CenterBold']))
    elements.append(Paragraph(f"Шартнома №{contract.code}", styles['CenterBold']))
    elements.append(Spacer(1, 12))

    date = Paragraph(f"Бухоро ш.", styles['LeftBold'])
    location = Paragraph(f"{contract.created_date.strftime('%d.%m.%Y')}", styles['RightBold'])

    # Create a table with one row and two columns
    date_location_table = Table([[date, location]], 
                            colWidths=[doc.width / 2.0, doc.width / 2.0])

    date_location_table.setStyle(TableStyle([
        # Align left cell to left, right cell to right
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        # Remove all paddings
        ('LEFTPADDING', (0, 0), (0, 0), 0),
        ('RIGHTPADDING', (1, 0), (1, 0), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        # Vertically top-align (optional)
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))

    elements.append(date_location_table)
    elements.append(Spacer(1, 10))



    elements.append(Paragraph(
        f"<font name='TimesNewRomanBold'>{contract.partner.department_name}</font> "
        f"(бундан буён - “Бажарувчи” деб аталади) номидан ҳаракат қилувчи директор "
        f"<font name='TimesNewRomanBold'>{contract.partner.head_of_company}</font>, бир томондан, ва "
        f"<font name='TimesNewRomanBold'>{contract.get_client_name()}</font> "
        f"(бундан буён – “Буюртмачи” деб аталади) номидан ҳаракат қилувчи директор "
        f"<font name='TimesNewRomanBold'>{contract.get_head_of_company()}</font>, иккинчи томондан,"
        f"биргаликда “Тарафлар”, алоҳида эса “Тараф” деб аталувчилар мазкур шартномани қуйидагилар ҳақида туздилар:",
        styles['Justify']
    ))
    elements.append(Spacer(1, 12))


    elements.append(Paragraph("1.ШАРТНОМА ПРЕДМЕТИ", styles['CenterBold']))
    elements.append(Spacer(1, 4))

    elements.append(Paragraph("<font name='TimesNewRomanBold'>Ушбу шартномага биноан, Мижоз тўлайди ва қабул қилади ва Ижрочи қуйидаги шартлар асосида товарларни (хизматларни) етказиб беради:</font>", styles['Justify']))
    elements.append(Spacer(1, 4))

    # Table Header
    table_data = [[
        Paragraph("№", styles['TinyCenteredBold']),
        Paragraph("Маҳсулот номи<br/>(хизматлар)", styles['TinyCenteredBold']),
        Paragraph("Товар (хизмат)лар Ягона<br/>электрон миллий каталоги бўйича<br/>идентификация коди ва номи", styles['TinyCenteredBold']),
        Paragraph("Товар/ хизмат штрих коди", styles['TinyCenteredBold']),
        Paragraph("Ўлчов бирлиг", styles['TinyCenteredBold']),
        Paragraph("Миқдори", styles['TinyCenteredBold']),
        Paragraph("Нарх", styles['TinyCenteredBold']),
        Paragraph("Етказиб бериш қиймати<br/>(ҚҚС сиз)", styles['TinyCenteredBold']),
        Paragraph("ҚҚС<br/>ставка", styles['TinyCenteredBold']),
        Paragraph("ҚҚС<br/>сумма", styles['TinyCenteredBold']),
        Paragraph("Етказиш баҳоси<br/>ҚҚС билан", styles['TinyCenteredBold']),
    ]]

    total_sum = Decimal("0")

    for i, s in enumerate(sales, start=1):
        total = s.cost * s.amount
        total_sum += total
        no_tax = total / Decimal("1.12")
        tax = total - no_tax
        if s.item.measurement == '1': 
            olchov = "Dona"
        else:
            olchov = "Litr"       
        table_data.append([
            Paragraph(str(i), styles['TinyCentered']),
            Paragraph(s.item.item_name, styles['TinyCentered']),
            Paragraph(s.item.code or "", styles['TinyCentered']),
            Paragraph(s.item.cypher_code or "", styles['TinyCentered']),
            Paragraph(olchov, styles['TinyCentered']),
            Paragraph(spaced_float(s.amount), styles['TinyCentered']),
            Paragraph(spaced_float(s.cost), styles['TinyCentered']),
            Paragraph(spaced_float(no_tax), styles['TinyCentered']),
            Paragraph("12%", styles['TinyCentered']),
            Paragraph(spaced_float(tax), styles['TinyCentered']),
            Paragraph(spaced_float(total), styles['TinyCentered']),
        ])

    # Final totals
    total_no_tax = total_sum / Decimal("1.12")
    total_tax = total_sum - total_no_tax

    table_data.append([
        "",
        Paragraph("<em>ЖАМИ:</em>", styles['TinyCenteredBold']),
        "", "", "", "", "",
        Paragraph(spaced_float(total_no_tax), styles['TinyCenteredBold']),
        Paragraph("", styles['TinyCenteredBold']),
        Paragraph(spaced_float(total_tax), styles['TinyCenteredBold']),
        Paragraph(spaced_float(total_sum), styles['TinyCenteredBold']),
    ])
    col_weights = [1, 4, 6, 2, 2, 2, 4, 4, 2, 4, 4]
    total_weight = sum(col_weights)
    col_widths = [(weight / total_weight) * doc.width for weight in col_weights]

    sales_table = Table(table_data, colWidths=col_widths, hAlign='LEFT', repeatRows=1)
    sales_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),

        # Horizontal alignment handled via Paragraph styles ('Centered', etc.)
        # Vertical alignment handled here:
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),

        # Optional: Padding control
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))


    elements.append(sales_table)
    elements.append(Spacer(1, 12))

    # Total in words (placeholder, implement num_to_uz_text if needed)
    elements.append(Paragraph(f"<font name='TimesNewRomanBold'>Шартноманинг умумий миқдори: ҚҚС инобатга олган холда - {num_to_uz_cyrillic_text(total_sum)} сўм</font> (сумма сўзлар билан), ({spaced_float(total_sum)})", styles['Centered']))
    elements.append(Spacer(1, 12))
        
    elements.append(Paragraph(
        "1.1. <font name='TimesNewRomanBold'>“СОТУВЧИ”</font> қуйидаги товарларни етказиб бериш, "
        "<font name='TimesNewRomanBold'>“СОТИБ ОЛУВЧИ”</font> товарлар учун шартномада кўрсатилган муддатларда хақ тўлаш "
        "ва товарларни тўлиқ миқдорда қабул қилиб олиш мажбуриятини ўз зиммасига олади."
        "<br/>1.2. Товар ҳақида маълумотлар Ушбу шартноманинг № 1-жадвалида кўрсатилган ёки товар-транспорт юк хати орқали юборилган бўлади."
        "<br/>1.3. <font name='TimesNewRomanBold'>“СОТИБ ОЛУВЧИ”</font> 1-жадвалда кўрсатилган жами шартнома суммасидан ошмаган ҳолда бошқа товарлар олиш хуқуқига эга.",
        styles['Justify']
    ))
    elements.append(Spacer(1, 8))

    custom_contexts = ContractContext.objects.filter(contract=contract, deleted=False)
    context_dict = {}
    for ctx in custom_contexts:
        # Key: (paragraph_number, key, value)
        context_dict[(ctx.paragraph_number, ctx.key, ctx.value)] = ctx.paragraph_text
    
    all_paragraph_numbers = sorted(contract.paragraph_numbers)
    
    for num in all_paragraph_numbers:
        defaults = get_default_paragraph_text(num)
        # 1. Get header (always value=1 for header)
        header = context_dict.get((num, 'header', 1), defaults['header'][1])
        # 2. Get all subparagraphs (text)
        subparagraphs = []
        for subkey, default_text in defaults['text'].items():
            sub_text = context_dict.get((num, 'text', subkey), default_text)
            # Clean for PDF, add to list
            subparagraphs.append(f"{num}.{subkey}. {clean_paragraph_for_pdf(sub_text)}")
        # Join all subparagraphs into one body
        body = "".join(subparagraphs)

        # Build PDF elements
        elements.append(Paragraph(f"<font name='TimesNewRomanBold'>{num}. {header}</font>", styles['CenterBold']))
        elements.append(Paragraph(body, styles['Justify']))
        elements.append(Spacer(1, 6))
    
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(
        f" {max(all_paragraph_numbers) + 1}. ТОМОНЛАРНИНГ ЮРИДИК МАНЗИЛИ ВА РЕКВИЗИТЛАРИ",
        styles['CenterBold']
    ))
    elements.append(Spacer(1, 6))
    
    # Footer Table
    seller_info = [
        Paragraph(f"<font name='TimesNewRomanBold'>{contract.partner.department_name}</font>", styles['CenteredBold']),
        Paragraph(f"<font name='TimesNewRomanBold'>Адрес:</font> {contract.partner.address}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>ИНН:</font> {contract.partner.tin_number}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>Хисоб рақам:</font> {contract.partner.account_number}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>Банк:</font> {contract.partner.bank_address}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>МФО:</font> {contract.partner.mfo}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>Тел:</font> {contract.partner.phone_number}", styles['Justify']),
    ]


    if contract.client_company:
        buyer_info = [
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_client_name()}</font>", styles['CenteredBold']),
            Paragraph(f"<font name='TimesNewRomanBold'>Адрес:</font> {contract.client_company.address}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Хисоб рақам:</font> {contract.client_company.account_number}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Банк:</font> {contract.client_company.bank_address}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>МФО:</font> {contract.client_company.mfo}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>ИНН:</font> {contract.get_tin_number()}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Тел:</font> {contract.get_phone_number()}", styles['Justify']),
        ]
    else:
        buyer_info = [
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_client_name()}</font>", styles['CenteredBold']),
            Paragraph(f"<font name='TimesNewRomanBold'>Адрес:</font> {contract.client_person.address}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>ЖШШИР:</font> {contract.client_person.pinfl}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Паспорт серия:</font> {contract.client_person.passport_number}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>ИНН:</font> {contract.get_tin_number()}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Тел:</font> {contract.get_phone_number()}", styles['Justify']),
        ]

    seller_cell = []
    for para in seller_info:
        seller_cell.append(para)
        seller_cell.append(Spacer(1, 4))  # optional spacing between lines
    buyer_cell = []
    
    for para in buyer_info:
        buyer_cell.append(para)
        buyer_cell.append(Spacer(1, 4))
    footer_table = Table([
        [
            Paragraph("<font name='TimesNewRomanBold'>СОТУВЧИ</font>", styles['CenterBold']),
            Paragraph("<font name='TimesNewRomanBold'>СОТИБ ОЛУВЧИ</font>", styles['CenterBold']),
        ],
        [
            seller_cell,
            buyer_cell
        ],
        [
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.partner.head_of_company}</font>", styles['CenterBold']),
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_head_of_company()}</font>", styles['CenterBold']),
        ],
    ], colWidths=[220, 220])

    footer_table.setStyle(TableStyle([
        # ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # header row center aligned
        ('VALIGN', (0, 1), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), 'TimesNewRoman'),  # default font fallback
        ('LEFTPADDING', (0, 0), (-1, -1), 15),
        ('RIGHTPADDING', (0, 0), (-1, -1), 15),
        #('TOPPADDING', (0, 0), (-1, -1), 4),
        #('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(footer_table)
    elements.append(Spacer(1, 20))
    if contract.is_public:
        elements.append(get_qr_image(contract.pk))


    doc.build(elements)

    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{contract.get_client_name()}_{contract.code}.pdf"'
     #If you want it to open in browser, use:
     #  response['Content-Disposition'] = f'inline; filename="contract_{contract.code}.pdf"'

    return response






def generate_contract_canceled_pdf(pk):
    contract = get_object_or_404(ContractNew, pk=pk, deleted=False)

    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'TIMES.TTF')
    pdfmetrics.registerFont(TTFont('TimesNewRoman', font_path))
    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'TIMESBD.TTF')
    pdfmetrics.registerFont(TTFont('TimesNewRomanBold', font_path))

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=25*mm, rightMargin=25*mm, topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY, fontName='TimesNewRoman', fontSize=11, leading=14))
    styles.add(ParagraphStyle(name='CenterBold', alignment=TA_CENTER, fontName='TimesNewRomanBold', fontSize=11))
    styles.add(ParagraphStyle(name='Centered', alignment=1, fontName='TimesNewRoman', fontSize=11, leading=11,))
    styles.add(ParagraphStyle(name='CenteredBold', parent=styles['Centered'], fontName='TimesNewRomanBold',))
    styles.add(ParagraphStyle(
        name='LeftBold',
        parent=styles['Normal'],
        fontName='TimesNewRomanBold',  # or 'TimesNewRomanBold' if registered
        alignment=0  # 0 = left
    ))
    styles.add(ParagraphStyle(
    name='RightBold',
    parent=styles['Normal'],
    fontName='TimesNewRomanBold',
    alignment=2  # 2 = right
    ))
    elements = []

    # Title
    elements.append(Paragraph(f"{contract.code} - ШАРТНОМАНИ БЕКОР ҚИЛИШ  КЕЛИШУВ БАЁННОМАСИ", styles['CenterBold']))
    elements.append(Spacer(1, 12))

    date = Paragraph(f"{contract.updated_at.strftime('%Y-%m-%d')}", styles['LeftBold'])
    location = Paragraph("Бухоро шахар", styles['RightBold'])

    # Create a table with one row and two columns
    date_location_table = Table([[date, location]], 
                            colWidths=[doc.width / 2.0, doc.width / 2.0])

    date_location_table.setStyle(TableStyle([
        # Align left cell to left, right cell to right
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        # Remove all paddings
        ('LEFTPADDING', (0, 0), (0, 0), 0),
        ('RIGHTPADDING', (1, 0), (1, 0), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        # Vertically top-align (optional)
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))

    elements.append(date_location_table)
    elements.append(Spacer(1, 10))



    elements.append(Paragraph(
        f"{contract.created_date.strftime('%Y-%m-%d')} даги "
        f"<font name='TimesNewRomanBold'>{contract.code}</font> - сонли олди-сотди шартномаси бўйича томонлар "
        f"<font name='TimesNewRomanBold'>«Сотувчи»</font> деб номланувчи <font name='TimesNewRomanBold'>{contract.partner.department_name}</font> низомига асосан иш юритувчи "
        f"<font name='TimesNewRomanBold'>директор {contract.partner.head_of_company}</font> бир томондан  "
        f"<font name='TimesNewRomanBold'>«Сотиб олувчи» {contract.get_client_name()}</font> низомига асосан иш юритувчи "
        f"<font name='TimesNewRomanBold'>{contract.get_head_of_company()}</font> иккинчи томондан шартномани бекор қилишга келишувга келдилар.",
        styles['Justify']
    ))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(
        f"Кўчирилган аванс пулини <font name='TimesNewRomanBold'>«Сотиб олувчи»</font>нинг қуйидаги хисоб рақамига кўчирилади:",
        styles['Justify']
    ))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_client_name()}</font>", styles['Justify']))
    elements.append(Paragraph(f"Х/Р: {contract.get_client().account_number}", styles['Justify']))
    elements.append(Paragraph(f"МФО: {contract.get_client().mfo}", styles['Justify']))
    elements.append(Paragraph(f"ИНН: {contract.get_client().tin_number}", styles['Justify']))

    elements.append(Spacer(1, 8))
    
    footer_table = Table([
        [
            Paragraph("<font name='TimesNewRomanBold'>СОТУВЧИ</font>", styles['CenterBold']),
            Paragraph("<font name='TimesNewRomanBold'>СОТИБ ОЛУВЧИ</font>", styles['CenterBold']),
        ],
        [
            Paragraph(f"", styles['CenterBold']),
            Paragraph(f"", styles['CenterBold']),
            
        ],
        [
            Paragraph(f"_______________", styles['CenterBold']),
            Paragraph(f"_______________", styles['CenterBold']),
            
        ],
        [
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.partner.head_of_company}</font>", styles['CenterBold']),
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_head_of_company()}</font>", styles['CenterBold']),
        ],
    ], colWidths=[220, 220])

    footer_table.setStyle(TableStyle([
        # ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # header row center aligned
        ('VALIGN', (0, 1), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), 'TimesNewRoman'),  # default font fallback
        ('LEFTPADDING', (0, 0), (-1, -1), 15),
        ('RIGHTPADDING', (0, 0), (-1, -1), 15),
        #('TOPPADDING', (0, 0), (-1, -1), 4),
        #('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(footer_table)
    elements.append(Spacer(1, 20))
    if contract.is_public:
        elements.append(get_qr_image(contract.pk))

    doc.build(elements)

    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{contract.get_client_name()}_{contract.code}.pdf"'
     #If you want it to open in browser, use:
     #  response['Content-Disposition'] = f'inline; filename="contract_{contract.code}.pdf"'

    return response


def get_qr_agreement_base64(pk, num):
    contract = get_object_or_404(ContractNew, pk=pk, is_public=True, deleted=False)
    url = f"https://roadstar.uz/contracts/public-pdf/{contract.public_token}/agreement/{num}"    

    qr = qrcode.make(url)
    buffer = BytesIO()
    qr.save(buffer, format="PNG")
    img_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
    return f"data:image/png;base64,{img_str}"

def get_qr_agreement(pk, num):
    contract = get_object_or_404(ContractNew, pk=pk, is_public=True, deleted=False)
    url = f"https://roadstar.uz/contracts/public-pdf/{contract.public_token}/agreement/{num}"

    qr = qrcode.make(url)
    buffer = BytesIO()
    qr.save(buffer, format="PNG")
    buffer.seek(0)
    return Image(buffer, width=80, height=80)


def generate_agreement_pdf(pk, num):
    contract = get_object_or_404(ContractNew, pk=pk, deleted=False)
    additional = Agreement.objects.filter(contract=contract, pk=num, deleted=False).first()
    context = AgreementContext.objects.filter(agreement_num=additional, deleted=False)
    
    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'TIMES.TTF')
    pdfmetrics.registerFont(TTFont('TimesNewRoman', font_path))
    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'TIMESBD.TTF')
    pdfmetrics.registerFont(TTFont('TimesNewRomanBold', font_path))

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=25*mm, rightMargin=25*mm, topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY, fontName='TimesNewRoman', fontSize=11, leading=14))
    styles.add(ParagraphStyle(name='CenterBold', alignment=TA_CENTER, fontName='TimesNewRomanBold', fontSize=11))
    styles.add(ParagraphStyle(name='Centered', alignment=1, fontName='TimesNewRoman', fontSize=11, leading=11,))
    styles.add(ParagraphStyle(name='CenteredBold', parent=styles['Centered'], fontName='TimesNewRomanBold',))
    styles.add(ParagraphStyle(
        name='LeftBold',
        parent=styles['Normal'],
        fontName='TimesNewRomanBold',  # or 'TimesNewRomanBold' if registered
        alignment=0  # 0 = left
    ))
    styles.add(ParagraphStyle(
    name='RightBold',
    parent=styles['Normal'],
    fontName='TimesNewRomanBold',
    alignment=2  # 2 = right
    ))
    elements = []

    # Title
    elements.append(Paragraph(f"{ contract.created_date.strftime('%d-%m-%Y') } ЙИЛДАГИ №{contract.code} ШАРТНОМАГА", styles['CenterBold']))
    elements.append(Paragraph(f"ҚЎШИМЧА КЕЛИШУВ №{ additional.code }", styles['CenterBold']))
    elements.append(Spacer(1, 12))

    date = Paragraph(f"{additional.created_date.strftime('%Y-%m-%d')}", styles['LeftBold'])
    location = Paragraph("Бухоро шахар", styles['RightBold'])

    # Create a table with one row and two columns
    date_location_table = Table([[date, location]], 
                            colWidths=[doc.width / 2.0, doc.width / 2.0])

    date_location_table.setStyle(TableStyle([
        # Align left cell to left, right cell to right
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        # Remove all paddings
        ('LEFTPADDING', (0, 0), (0, 0), 0),
        ('RIGHTPADDING', (1, 0), (1, 0), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        # Vertically top-align (optional)
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))

    elements.append(date_location_table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph(
        f"<font name='TimesNewRomanBold'>{ contract.partner.department_name }</font> (кейинги ўринларда «Сотувчи» деб номланади) номидан Устав асосида фаолият юритувчи директор "
        f"<font name='TimesNewRomanBold'>{ contract.partner.head_of_company }</font> бир тарафдан ва "
        f"<font name='TimesNewRomanBold'>{ contract.get_client_name() }</font> "
        f"(кейинги ўринларда «Сотиб олувчи» деб номланади) устав асосида фаолият юритувчи раҳбар "
        f"<font name='TimesNewRomanBold'>{ contract.get_head_of_company() }</font> иккинчи тарафдан, мазкур қўшимча келишувни қуйидагилар хақида туздилар: ",
        styles['Justify']
    ))
    elements.append(Spacer(1, 10))

    subs = ''
    if context:
        subs = ''.join(add.paragraph_number for add in context)

    elements.append(Paragraph(f"1.ШАРТНОМА МАЗМУНИ.", styles['CenterBold']))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(
            f"<font name='TimesNewRomanBold'>Мазкур қўшичмча келишув</font> "
            f"{contract.created_date.strftime('%Y-%m-%d')} йилдаги "
            f"{contract.code} шартноманинг {subs} пунктларига ва реквизитларига "
            f"ўзгартириш киритиш тўғрисида тузилди.",
        styles['Justify']
    ))
    elements.append(Spacer(1, 8))
    
    if context:
        for add in context:
            elements.append(Paragraph(f"{add.paragraph_number}. { clean_paragraph_for_pdf(add.paragraph_text) }", styles['Justify']))
    else:
        elements.append(Paragraph(f" ... ", styles['Justify']))
    elements.append(Spacer(1, 12))
    
    elements.append(Paragraph(
        f"ТОМОНЛАРНИНГ ЮРИДИК МАНЗИЛИ ВА РЕКВИЗИТЛАРИ",
        styles['CenterBold']
    ))
    elements.append(Spacer(1, 6))
    
    # Footer Table
    seller_info = [
        Paragraph(f"<font name='TimesNewRomanBold'>{contract.partner.department_name}</font>", styles['CenteredBold']),
        Paragraph(f"<font name='TimesNewRomanBold'>Адрес:</font> {contract.partner.address}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>ИНН:</font> {contract.partner.tin_number}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>Хисоб рақам:</font> {contract.partner.account_number}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>Банк:</font> {contract.partner.bank_address}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>МФО:</font> {contract.partner.mfo}", styles['Justify']),
        Paragraph(f"<font name='TimesNewRomanBold'>Тел:</font> {contract.partner.phone_number}", styles['Justify']),
    ]


    if contract.client_company:
        buyer_info = [
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_client_name()}</font>", styles['CenteredBold']),
            Paragraph(f"<font name='TimesNewRomanBold'>Адрес:</font> {contract.client_company.address}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Хисоб рақам:</font> {contract.client_company.account_number}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Банк:</font> {contract.client_company.bank_address}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>МФО:</font> {contract.client_company.mfo}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>ИНН:</font> {contract.get_tin_number()}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Тел:</font> {contract.get_phone_number()}", styles['Justify']),
        ]
    else:
        buyer_info = [
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_client_name()}</font>", styles['CenteredBold']),
            Paragraph(f"<font name='TimesNewRomanBold'>Адрес:</font> {contract.client_person.address}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>ЖШШИР:</font> {contract.client_person.pinfl}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Паспорт серия:</font> {contract.client_person.passport_number}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>ИНН:</font> {contract.get_tin_number()}", styles['Justify']),
            Paragraph(f"<font name='TimesNewRomanBold'>Тел:</font> {contract.get_phone_number()}", styles['Justify']),
        ]

    seller_cell = []
    for para in seller_info:
        seller_cell.append(para)
        seller_cell.append(Spacer(1, 4))  # optional spacing between lines
    
    buyer_cell = []
    for para in buyer_info:
        buyer_cell.append(para)
        buyer_cell.append(Spacer(1, 4))
    
    footer_table = Table([
        [
            Paragraph("<font name='TimesNewRomanBold'>СОТУВЧИ</font>", styles['CenterBold']),
            Paragraph("<font name='TimesNewRomanBold'>СОТИБ ОЛУВЧИ</font>", styles['CenterBold']),
        ],
        [
            seller_cell,
            buyer_cell
        ],
        [
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.partner.head_of_company}</font>", styles['CenterBold']),
            Paragraph(f"<font name='TimesNewRomanBold'>{contract.get_head_of_company()}</font>", styles['CenterBold']),
        ],
    ], colWidths=[220, 220])

    footer_table.setStyle(TableStyle([
        # ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # header row center aligned
        ('VALIGN', (0, 1), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), 'TimesNewRoman'),  # default font fallback
        ('LEFTPADDING', (0, 0), (-1, -1), 15),
        ('RIGHTPADDING', (0, 0), (-1, -1), 15),
        #('TOPPADDING', (0, 0), (-1, -1), 4),
        #('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(footer_table)
    elements.append(Spacer(1, 20))
    
    if additional.is_public:
        elements.append(get_qr_agreement(contract.pk, additional.pk))


    doc.build(elements)

    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{contract.get_client_name()}_{contract.code}_{additional.pk}.pdf"'
     #If you want it to open in browser, use:
     #  response['Content-Disposition'] = f'inline; filename="contract_{contract.code}.pdf"'

    return response



import docx
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL

def populate_paragraph_from_html(paragraph, html_text):
    """
    Populates a given paragraph object with text parsed from a simple HTML string
    containing <b> and <br/> tags.

    :param paragraph: The python-docx paragraph object to populate.
    :param html_text: The string containing the text with simple HTML tags.
    """
    # Sanitize the input to be a string
    html_text = html_text or ''

    # Split the entire text by the <br/> tag to handle line breaks
    segments = html_text.split('<br/>')

    for i, segment in enumerate(segments):
        # Add a soft line break for every segment after the first one
        if i > 0:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.LINE)

        # Use regex to split the segment by <b> tags, keeping the tags
        # This will give us a list of alternating normal and bold text parts
        parts = re.split(r'(<b>.*?</b>)', segment)
        
        for part in parts:
            # Skip any empty strings that might result from the split
            if not part:
                continue

            if part.startswith('<b>') and part.endswith('</b>'):
                # If the part is a <b> tag, strip the tags and add a bold run
                content = part[3:-4]  # Remove <b> and </b>
                paragraph.add_run(content).bold = True
            else:
                # Otherwise, add a normal run
                paragraph.add_run(part)



def generate_contract_docx(pk):
    """
    Generates a DOCX contract file based on the ContractNew object.
    """
    contract = get_object_or_404(ContractNew, pk=pk, deleted=False)
    sales = ContractItem.objects.filter(deleted=False, contract_id=contract.pk).exclude(item='1').select_related('item')

    # --- 1. Document Setup ---
    document = docx.Document()

    # Set default font for the entire document
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # Set page margins
    section = document.sections[0]
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    # --- 2. Title and Header ---
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"ОЛДИ-СОТДИ ШАРТНОМАСИ\nШартнома №{contract.code}")
    title_run.bold = True
    
    document.add_paragraph() # Spacer

    # Using a 1x2 table for left/right aligned text
    header_table = document.add_table(rows=1, cols=2)
    cell_left = header_table.cell(0, 0)
    cell_right = header_table.cell(0, 1)

    p_left = cell_left.paragraphs[0]
    p_left.add_run("Бухоро ш.").bold = True
    
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.add_run(f"{contract.created_date.strftime('%d.%m.%Y')}").bold = True
    
    # --- 3. Introductory Paragraph ---
    intro_p = document.add_paragraph()
    intro_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro_p.add_run(f"{contract.partner.department_name}").bold = True
    intro_p.add_run(" (бундан буён - “Бажарувчи” деб аталади) номидан ҳаракат қилувчи директор ")
    intro_p.add_run(f"{contract.partner.head_of_company}").bold = True
    intro_p.add_run(", бир томондан, ва ")
    intro_p.add_run(f"{contract.get_client_name()}").bold = True
    intro_p.add_run(" (бундан буён – “Буюртмачи” деб аталади) номидан ҳаракат қилувчи директор ")
    intro_p.add_run(f"{contract.get_head_of_company()}").bold = True
    intro_p.add_run(", иккинчи томондан, биргаликда “Тарафлар”, алоҳида эса “Тараф” деб аталувчилар мазкур шартномани қуйидагилар ҳақида туздилар:")

    document.add_paragraph() # Spacer

    # --- 4. Subject of Contract & Sales Table ---
    p_subject_header = document.add_paragraph()
    p_subject_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subject_header.add_run("1.ШАРТНОМА ПРЕДМЕТИ").bold = True
    
    p_subject_intro = document.add_paragraph()
    p_subject_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_subject_intro.add_run("Ушбу шартномага биноан, Мижоз тўлайди ва қабул қилади ва Ижрочи қуйидаги шартлар асосида товарларни (хизматларни) етказиб беради:").bold = True

    # Sales Table
    table_headers = ["№", "Товар номи", "1 донаси нархи\nҚҚС билан", "Миқдори", "Жами суммаси\nҚҚС сисиз", "ҚҚС суммаси", "Жами суммаси\nҚҚС билан"]
    sales_table = document.add_table(rows=1, cols=len(table_headers))
    sales_table.style = 'Table Grid'
    hdr_cells = sales_table.rows[0].cells
    for i, header_text in enumerate(table_headers):
        p = hdr_cells[i].paragraphs[0]
        p.add_run(header_text).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    total_sum = Decimal("0")
    for i, s in enumerate(sales, start=1):
        total = s.cost * s.amount
        total_sum += total
        no_tax = total / Decimal("1.12")
        tax = total - no_tax
        
        row_data = [str(i), s.item.item_name, f"{s.cost:,.2f}", str(s.amount), f"{no_tax:,.2f}", f"{tax:,.2f}", f"{total:,.2f}"]
        row_cells = sales_table.add_row().cells
        for j, cell_text in enumerate(row_data):
            p = row_cells[j].paragraphs[0]
            p.add_run(cell_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Totals Row
    total_no_tax = total_sum / Decimal("1.12")
    total_tax = total_sum - total_no_tax
    total_cells = sales_table.add_row().cells
    total_cells[1].paragraphs[0].add_run("ЖАМИ:").bold = True
    total_cells[4].paragraphs[0].add_run(f"{total_no_tax:,.2f}").bold = True
    total_cells[5].paragraphs[0].add_run(f"{total_tax:,.2f}").bold = True
    total_cells[6].paragraphs[0].add_run(f"{total_sum:,.2f}").bold = True
    for cell in total_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    document.add_paragraph() # Spacer

    # Total in words
    total_p = document.add_paragraph()
    total_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_p.add_run("Жами:").bold = True
    total_p.add_run(f" ҚҚС инобатга олган холда - ")
    # Assuming num_to_uz_cyrillic_text returns a string
    total_p.add_run(f"{total_sum:,.2f} ({num_to_uz_cyrillic_text(total_sum)}) сўм.").bold = True

    document.add_paragraph() # Spacer

    # Static paragraphs after the table
    p1_1 = document.add_paragraph()
    p1_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1_1.add_run("1.1. “СОТУВЧИ”").bold = True
    p1_1.add_run(" қуйидаги товарларни етказиб бериш, ")
    p1_1.add_run("“СОТИБ ОЛУВЧИ”").bold = True
    p1_1.add_run(" товарлар учун шартномада кўрсатилган муддатларда хақ тўлаш ва товарларни тўлиқ миқдорда қабул қилиб олиш мажбуриятини ўз зиммасига олади.\n1.2. Товар ҳақида маълумотлар Ушбу шартноманинг № 1-жадвалида кўрсатилган ёки товар-транспорт юк хати орқали юборилган бўлади.\n1.3. ")
    p1_1.add_run("“СОТИБ ОЛУВЧИ”").bold = True
    p1_1.add_run(" 1-жадвалда кўрсатилган жами шартнома суммасидан ошмаган ҳолда бошқа товарлар олиш хуқуқига эга.")

    # --- 5. Dynamic Paragraphs ---
    # This logic remains largely the same
    custom_contexts = ContractContext.objects.filter(contract=contract, deleted=False)
    context_dict = {(ctx.paragraph_number, ctx.key, ctx.value): ctx.paragraph_text for ctx in custom_contexts}
    all_paragraph_numbers = sorted(contract.paragraph_numbers)

    for num in all_paragraph_numbers:
        defaults = get_default_paragraph_text(num)
        header = context_dict.get((num, 'header', 1), defaults['header'][1])
        
        subparagraphs = []
        for subkey, default_text in defaults['text'].items():
            sub_text = context_dict.get((num, 'text', subkey), default_text)
            subparagraphs.append(f"{num}.{subkey}. {sub_text}")
        body = "".join(subparagraphs)

        p_header = document.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_header.add_run(f"{num}. {header}").bold = True

        p_body = document.add_paragraph()  # Create an empty paragraph
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        populate_paragraph_from_html(p_body, body) # Use the helper to fill it

    # --- 6. Footer with Party Details ---
    document.add_paragraph() # Spacer
    footer_header = document.add_paragraph()
    footer_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_header.add_run(f"{max(all_paragraph_numbers) + 1}. ТОМОНЛАРНИНГ ЮРИДИК МАНЗИЛИ ВА РЕКВИЗИТЛАРИ").bold = True
    
    footer_table = document.add_table(rows=1, cols=2)
    seller_cell = footer_table.cell(0, 0)
    buyer_cell = footer_table.cell(0, 1)

    # Seller Info
    seller_p = seller_cell.add_paragraph()
    seller_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    seller_p.add_run("СОТУВЧИ\n").bold = True
    seller_p.add_run(f"{contract.partner.department_name}").bold = True
    
    seller_details = seller_cell.add_paragraph()
    seller_details.add_run("Адрес: ").bold = True
    seller_details.add_run(f"{contract.partner.address}\n")
    seller_details.add_run("ИНН: ").bold = True
    seller_details.add_run(f"{contract.partner.tin_number}\n")
    seller_details.add_run("Хисоб рақам: ").bold = True
    seller_details.add_run(f"{contract.partner.account_number}\n")
    seller_details.add_run("Банк: ").bold = True
    seller_details.add_run(f"{contract.partner.bank_address}\n")
    seller_details.add_run("МФО: ").bold = True
    seller_details.add_run(f"{contract.partner.mfo}\n")  
    seller_details.add_run("Тел: ").bold = True
    seller_details.add_run(f"{contract.partner.phone_number}")

    seller_head = seller_cell.add_paragraph(f"\n\n{contract.partner.head_of_company}")
    seller_head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Buyer Info
    buyer_p = buyer_cell.add_paragraph()
    buyer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buyer_p.add_run("СОТИБ ОЛУВЧИ\n").bold = True
    buyer_p.add_run(f"{contract.get_client_name()}").bold = True

    buyer_details = buyer_cell.add_paragraph()
    if contract.client_company:
        buyer_details.add_run("Адрес: ").bold = True
        buyer_details.add_run(f"{contract.client_company.address}\n")
        buyer_details.add_run("Хисоб рақам: ").bold = True
        buyer_details.add_run(f"{contract.client_company.account_number}\n")
        buyer_details.add_run("Банк: ").bold = True
        buyer_details.add_run(f"{contract.client_company.bank_address}\n")
        buyer_details.add_run("МФО: ").bold = True
        buyer_details.add_run(f"{contract.client_company.mfo}\n")
    else: 
        buyer_details.add_run("Адрес: ").bold = True
        buyer_details.add_run(f"{contract.client_person.address}\n")
        buyer_details.add_run("ЖШШИР: ").bold = True
        buyer_details.add_run(f"{contract.client_person.pinfl}\n")
        buyer_details.add_run("Паспорт серия: ").bold = True
        buyer_details.add_run(f"{contract.client_person.passport_number}\n")

    buyer_details.add_run("ИНН: ").bold = True
    buyer_details.add_run(f"{contract.get_tin_number()}\n")
    buyer_details.add_run("Тел: ").bold = True
    buyer_details.add_run(f"{contract.get_phone_number()}")

    buyer_head = buyer_cell.add_paragraph(f"\n\n{contract.get_head_of_company()}")
    buyer_head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 8. Save to buffer and return response ---
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{contract.get_client_name()}_{contract.code}.docx"'
    return response